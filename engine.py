# ============================================================
# engine.py — Worksheet Generator core engine
# Ported from Worksheet_Generator_Complete.ipynb into a reusable,
# stateless Python module for the Flask web app.
# ============================================================

import base64
import io
import random
import re

import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML as WeasyHTML

# ------------------------------------------------------------
# Urdu helpers
# ------------------------------------------------------------

def fix_urdu(text):
    if text is None:
        return ""
    text = str(text)
    if re.search(r'[\u0600-\u06FF]', text):
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    return text


def is_urdu_text(text):
    return bool(re.search(r'[\u0600-\u06FF]', str(text or "")))


LABELS = {
    "en": {
        "name": "Name", "class": "Class", "roll_no": "Roll No", "date": "Date",
        "answer_key": "Answer Key", "grade": "Grade", "subject": "Subject",
        "question": "Question", "correct_answer": "Correct Answer",
        "model_answer": "Model Answer", "teacher_copy": "Teacher Copy — Answer Key",
    },
    "ur": {
        "name": "نام", "class": "جماعت", "roll_no": "رول نمبر", "date": "تاریخ",
        "answer_key": "جوابات", "grade": "جماعت", "subject": "مضمون",
        "question": "سوال", "correct_answer": "درست جواب",
        "model_answer": "نمونہ جواب", "teacher_copy": "اساتذہ کے لیے — جوابات",
    },
}


def L(key, lang="en"):
    label = LABELS.get(lang, LABELS["en"]).get(key, key)
    return fix_urdu(label) if lang == "ur" else label


# ------------------------------------------------------------
# Theme engine
# ------------------------------------------------------------

def get_theme_for_grade(grade, forced_mode="auto"):
    if forced_mode in ("colorful", "professional"):
        return forced_mode
    return "colorful" if int(grade) <= 5 else "professional"


def base_css(primary, accent):
    return f"""
    * {{ box-sizing: border-box; }}
    body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 0; color: #2b2b2b; }}
    .urdu-text {{ font-family: 'Noto Nastaliq Urdu', 'Segoe UI', serif; direction: rtl; text-align: right; font-size: 16px; line-height: 2; }}
    .page {{ page-break-after: always; padding: 20px 28px; position: relative; min-height: 100%; }}
    .page:last-child {{ page-break-after: auto; }}
    .header-bar {{ display: flex; align-items: center; justify-content: space-between; border-bottom: 3px solid {primary}; padding-bottom: 8px; margin-bottom: 14px; }}
    .header-left {{ display: flex; align-items: center; gap: 10px; }}
    .school-name {{ font-size: 18px; font-weight: 700; color: {primary}; }}
    .logo-img {{ object-fit: contain; }}
    .student-info-bar {{ display: flex; flex-wrap: wrap; gap: 18px; background: #f4f7fb; border: 1px dashed {primary}; border-radius: 8px; padding: 8px 14px; margin-bottom: 16px; font-size: 13px; }}
    .student-info-bar span b {{ color: {primary}; }}
    .question-block {{ margin-bottom: 18px; padding-bottom: 10px; }}
    .q-number {{ font-weight: 700; color: {primary}; margin-right: 6px; }}
    .options-row {{ display: flex; flex-wrap: wrap; gap: 14px; margin-top: 6px; margin-left: 22px; }}
    .option-item {{ border: 1px solid #ccc; border-radius: 6px; padding: 4px 10px; font-size: 13px; }}
    .fill-blank-line {{ display: inline-block; border-bottom: 1.5px solid #333; min-width: 120px; margin: 0 4px; }}
    .match-column-table {{ width: 90%; margin: 8px auto 0 22px; border-collapse: collapse; }}
    .match-column-table td {{ border: 1px solid #bbb; padding: 6px 10px; font-size: 13px; }}
    .answer-lines .line {{ border-bottom: 1px solid #999; height: 22px; margin: 0 22px 4px 22px; }}
    .footer-bar {{ position: absolute; bottom: 10px; left: 28px; right: 28px; display: flex; justify-content: space-between; font-size: 11px; color: #777; border-top: 1px solid #ddd; padding-top: 4px; }}
    .answer-key-page {{ background: #fbfbfb; }}
    .answer-key-list .answer-key-item {{ padding: 6px 0; border-bottom: 1px dashed #ccc; font-size: 14px; }}
    """


def theme_extra_css(theme, accent):
    if theme == "colorful":
        return f"""
        body {{ background: #fffdf6; }}
        .header-bar {{ background: linear-gradient(90deg, {accent}33, transparent); border-radius: 10px; }}
        .question-block {{ background: #ffffff; border: 2px solid {accent}; border-radius: 14px; padding: 12px 16px; }}
        .q-number {{ font-size: 16px; }}
        .option-item {{ background: #fff8e1; border-color: {accent}; border-radius: 14px; }}
        """
    return """
    body { background: #ffffff; }
    .question-block { border-bottom: 1px solid #e2e2e2; }
    .option-item { background: #fafafa; border-radius: 4px; }
    """


GRADE_ICON_MAP = {
    "apple": "🍎", "banana": "🍌", "cat": "🐱", "dog": "🐶", "monkey": "🐵",
    "elephant": "🐘", "lion": "🦁", "fish": "🐟", "bird": "🐦", "sun": "☀️",
    "moon": "🌙", "star": "⭐", "flower": "🌸", "tree": "🌳", "car": "🚗",
    "ball": "⚽", "book": "📕", "house": "🏠", "cow": "🐄", "chicken": "🐔",
}


def add_icon_if_available(word):
    return GRADE_ICON_MAP.get((word or "").strip().lower(), "")


# ------------------------------------------------------------
# Cover page + header/footer
# ------------------------------------------------------------

def build_logo_html(config, size_px=None):
    if not config.get("school_logo_base64"):
        return ""
    size = size_px or config.get("logo_size_px", 70)
    return f'<img class="logo-img" src="{config["school_logo_base64"]}" style="width:{size}px;height:{size}px;">'


def build_cover_page_html(config, worksheet_title, grade, subject, style=None, lang="en"):
    style = style or config.get("cover_page_style", "modern")
    primary = config.get("primary_color", "#2E86AB")
    accent = config.get("accent_color", "#F6C90E")
    logo = build_logo_html(config, size_px=130)
    school_name = config.get("school_name", "")
    title_disp = fix_urdu(worksheet_title) if lang == "ur" else worksheet_title
    grade_lbl, subj_lbl = L("grade", lang), L("subject", lang)
    title_css = "urdu-text" if lang == "ur" else ""

    if style == "modern":
        body = f"""
        <div style="text-align:center; padding-top:80px;">
            {logo}
            <h1 style="color:{primary}; font-size:34px; margin-top:20px;">{school_name}</h1>
            <div style="width:120px; height:4px; background:{accent}; margin:14px auto;"></div>
            <h2 class="{title_css}" style="color:#333; font-size:26px;">{title_disp}</h2>
            <p style="font-size:18px; color:#555;">{grade_lbl}: {grade} &nbsp;|&nbsp; {subj_lbl}: {subject}</p>
        </div>
        """
    elif style == "classic":
        body = f"""
        <div style="text-align:center; border:6px double {primary}; margin:40px; padding:60px 30px;">
            {logo}
            <h1 style="font-family:Georgia, serif; color:{primary};">{school_name}</h1>
            <h2 class="{title_css}" style="font-family:Georgia, serif;">{title_disp}</h2>
            <p>{grade_lbl}: {grade} | {subj_lbl}: {subject}</p>
        </div>
        """
    else:
        body = f"""
        <div style="padding-top:100px; padding-left:60px;">
            {logo}
            <h1 style="color:{primary}; font-weight:300;">{school_name}</h1>
            <h3 class="{title_css}" style="color:#666; font-weight:300;">{title_disp} — {grade_lbl} {grade}, {subject}</h3>
        </div>
        """
    return f'<div class="page" style="page-break-after: always;">{body}</div>'


def build_student_info_bar(lang="en"):
    css_class = "urdu-text" if lang == "ur" else ""
    return f"""
    <div class="student-info-bar {css_class}">
        <span><b>{L('name', lang)}:</b> _____________________</span>
        <span><b>{L('class', lang)}:</b> _______</span>
        <span><b>{L('roll_no', lang)}:</b> _______</span>
        <span><b>{L('date', lang)}:</b> _______</span>
    </div>
    """


def build_page_header(config, worksheet_title, lang="en"):
    title_class = "urdu-text" if lang == "ur" else ""
    return f"""
    <div class="header-bar">
        <div class="header-left">
            {build_logo_html(config)}
            <div class="school-name">{config.get("school_name","")}</div>
        </div>
        <div class="{title_class}" style="font-size:14px; font-weight:600;">{fix_urdu(worksheet_title) if lang=='ur' else worksheet_title}</div>
    </div>
    {build_student_info_bar(lang)}
    """


def build_page_footer(config, page_num, total_pages, lang="en"):
    page_word = "صفحہ" if lang == "ur" else "Page"
    of_word = "از" if lang == "ur" else "of"
    page_word = fix_urdu(page_word) if lang == "ur" else page_word
    return f"""
    <div class="footer-bar">
        <span>{config.get("school_name","")}</span>
        <span>{page_word} {page_num} {of_word} {total_pages}</span>
    </div>
    """


# ------------------------------------------------------------
# Question renderers
# ------------------------------------------------------------

def _q_text(q, q_num, theme):
    text = q.get("text", "")
    urdu = is_urdu_text(text)
    css_class = "urdu-text" if urdu else ""
    display_text = fix_urdu(text) if urdu else text
    icon = ""
    if theme == "colorful" and text:
        icon = add_icon_if_available(text.split()[-1])
        icon = f" {icon}" if icon else ""
    return f'<div><span class="q-number">{q_num}.</span><span class="{css_class}">{display_text}{icon}</span></div>'


def render_mcq(q, q_num, theme):
    header = _q_text(q, q_num, theme)
    options_html = "".join(
        f'<div class="option-item">{chr(97+i)}) {opt}</div>'
        for i, opt in enumerate(q.get("options", []))
    )
    return f'<div class="question-block">{header}<div class="options-row">{options_html}</div></div>'


def render_fill_blank(q, q_num, theme):
    header = _q_text(q, q_num, theme)
    blanks = q.get("blanks", 1)
    blank_html = " ".join('<span class="fill-blank-line">&nbsp;</span>' for _ in range(blanks))
    return f'<div class="question-block">{header}<div style="margin-top:6px; margin-left:22px;">{blank_html}</div></div>'


def render_match_column(q, q_num, theme):
    header = _q_text(q, q_num, theme)
    left = q.get("left", [])
    right = q.get("right", [])
    if "_shuffled_right" not in q:
        right_shuffled = right[:]
        random.shuffle(right_shuffled)
        q["_shuffled_right"] = right_shuffled
    right_shuffled = q["_shuffled_right"]
    rows = "".join(
        f"<tr><td>{l}</td><td style='width:40px;text-align:center;'>—</td><td>{r}</td></tr>"
        for l, r in zip(left, right_shuffled)
    )
    table = f'<table class="match-column-table">{rows}</table>'
    return f'<div class="question-block">{header}{table}</div>'


def render_short_or_detailed(q, q_num, theme):
    header = _q_text(q, q_num, theme)
    lines = q.get("lines", 3 if q.get("type") == "short" else 6)
    lines_html = "".join('<div class="line"></div>' for _ in range(lines))
    return f'<div class="question-block">{header}<div class="answer-lines" style="margin-top:8px;">{lines_html}</div></div>'


RENDERERS = {
    "mcq": render_mcq,
    "fill_blank": render_fill_blank,
    "match_column": render_match_column,
    "short": render_short_or_detailed,
    "detailed": render_short_or_detailed,
}


def render_question(q, q_num, theme):
    renderer = RENDERERS.get(q.get("type"))
    if not renderer:
        return f'<div class="question-block">Unknown question type: {q.get("type")}</div>'
    return renderer(q, q_num, theme)


# ------------------------------------------------------------
# Answer key renderers
# ------------------------------------------------------------

def _answer_line(q_num, content, lang="en"):
    css = "urdu-text" if lang == "ur" else ""
    return f'<div class="answer-key-item"><span class="q-number">{q_num}.</span> <span class="{css}">{content}</span></div>'


def render_answer_mcq(q, q_num, lang):
    ans = q.get("answer", "—")
    ans_disp = fix_urdu(ans) if lang == "ur" else ans
    return _answer_line(q_num, ans_disp, lang)


def render_answer_fill_blank(q, q_num, lang):
    ans = q.get("answer", "—")
    if isinstance(ans, list):
        ans = ", ".join(ans)
    ans_disp = fix_urdu(ans) if lang == "ur" else ans
    return _answer_line(q_num, ans_disp, lang)


def render_answer_match_column(q, q_num, lang):
    left = q.get("left", [])
    right = q.get("right", [])
    pairs = ", ".join(f"{l} → {r}" for l, r in zip(left, right))
    pairs_disp = fix_urdu(pairs) if lang == "ur" else pairs
    return _answer_line(q_num, pairs_disp, lang)


def render_answer_short_detailed(q, q_num, lang):
    ans = q.get("answer", "—")
    ans_disp = fix_urdu(ans) if lang == "ur" else ans
    return _answer_line(q_num, ans_disp, lang)


ANSWER_RENDERERS = {
    "mcq": render_answer_mcq,
    "fill_blank": render_answer_fill_blank,
    "match_column": render_answer_match_column,
    "short": render_answer_short_detailed,
    "detailed": render_answer_short_detailed,
}


def render_answer_entry(q, q_num, lang="en"):
    renderer = ANSWER_RENDERERS.get(q.get("type"))
    if not renderer:
        return ""
    return renderer(q, q_num, lang)


def build_answer_key_html(config, questions, worksheet_title, lang="en", as_teacher_cover=True):
    title_lbl = L("answer_key", lang)
    title_disp = fix_urdu(worksheet_title) if lang == "ur" else worksheet_title
    title_css = "urdu-text" if lang == "ur" else ""

    items = "".join(render_answer_entry(q, i, lang) for i, q in enumerate(questions, start=1))

    cover = ""
    if as_teacher_cover:
        cover = f"""
        <div style="text-align:center; margin-bottom:20px; border-bottom:3px solid {config.get('primary_color','#2E86AB')}; padding-bottom:10px;">
            <h2 class="{title_css}">{fix_urdu(L('teacher_copy', lang)) if lang=='ur' else L('teacher_copy', lang)}</h2>
            <p class="{title_css}">{title_disp}</p>
        </div>
        """

    return f"""
    <div class="page answer-key-page">
        {cover}
        <h3 class="{title_css}">{fix_urdu(title_lbl) if lang=='ur' else title_lbl}</h3>
        <div class="answer-key-list">{items}</div>
    </div>
    """


# ------------------------------------------------------------
# Layout engine
# ------------------------------------------------------------

def randomize_question_order(questions, mode="same"):
    if mode == "different":
        shuffled = questions[:]
        random.shuffle(shuffled)
        return shuffled
    return questions


def paginate(questions, questions_per_page):
    if questions_per_page <= 0:
        return [questions]
    return [questions[i:i + questions_per_page] for i in range(0, len(questions), questions_per_page)]


def build_worksheet_pages_html(config, questions, worksheet_title, grade, theme,
                                questions_per_page=0, question_spacing_px=18,
                                binding_gap_px=0, lang="en"):
    pages = paginate(questions, questions_per_page)
    total_pages = len(pages)
    html_pages = []

    for idx, page_questions in enumerate(pages, start=1):
        q_html = ""
        counter_start = sum(len(p) for p in pages[:idx - 1]) + 1
        for offset, q in enumerate(page_questions):
            q_num = counter_start + offset
            q_html += render_question(q, q_num, theme)
            q_html += f'<div style="height:{question_spacing_px}px;"></div>'

        page_html = f"""
        <div class="page" style="padding-left:{28 + binding_gap_px}px;">
            {build_page_header(config, worksheet_title, lang)}
            <div class="questions-area">{q_html}</div>
            {build_page_footer(config, idx, total_pages, lang)}
        </div>
        """
        html_pages.append(page_html)

    return "".join(html_pages)


def build_full_worksheet_html(config, questions, worksheet_title, grade, subject,
                               theme=None, questions_per_page=0, question_spacing_px=18,
                               num_copies=1, randomize_mode="same", binding_gap_px=0,
                               cover_style=None, lang="en", answer_key_mode="separate"):
    theme = theme or get_theme_for_grade(grade, config.get("theme_mode", "auto"))
    primary = config.get("primary_color", "#2E86AB")
    accent = config.get("accent_color", "#F6C90E")
    css = base_css(primary, accent) + theme_extra_css(theme, accent)

    body_parts = []
    last_copy_questions = None

    for _ in range(max(1, num_copies)):
        copy_questions = randomize_question_order(questions, randomize_mode)
        last_copy_questions = copy_questions

        cover_html = build_cover_page_html(config, worksheet_title, grade, subject, style=cover_style, lang=lang)
        content_html = build_worksheet_pages_html(
            config, copy_questions, worksheet_title, grade, theme,
            questions_per_page=questions_per_page,
            question_spacing_px=question_spacing_px,
            binding_gap_px=binding_gap_px,
            lang=lang
        )
        body_parts.append(cover_html + content_html)

        if answer_key_mode in ("inline", "both"):
            body_parts.append(build_answer_key_html(config, copy_questions, worksheet_title, lang=lang, as_teacher_cover=False))

    student_html = f"""
    <html><head><meta charset="utf-8"><style>{css}</style></head>
    <body>{''.join(body_parts)}</body></html>
    """

    teacher_html = None
    if answer_key_mode in ("separate", "both"):
        answer_section = build_answer_key_html(config, last_copy_questions, worksheet_title, lang=lang, as_teacher_cover=True)
        teacher_html = f"""
        <html><head><meta charset="utf-8"><style>{css}</style></head>
        <body>{answer_section}</body></html>
        """

    return {"student_html": student_html, "teacher_html": teacher_html}


# ------------------------------------------------------------
# Export engine — PDF (WeasyPrint) + Word (python-docx)
# ------------------------------------------------------------

def export_to_pdf(full_html, filepath):
    WeasyHTML(string=full_html).write_pdf(filepath)
    return filepath


def export_worksheet_pdf(html_dict, base_filepath):
    student_file = f"{base_filepath}.pdf"
    export_to_pdf(html_dict["student_html"], student_file)

    teacher_file = None
    if html_dict.get("teacher_html"):
        teacher_file = f"{base_filepath}_ANSWER_KEY.pdf"
        export_to_pdf(html_dict["teacher_html"], teacher_file)

    return {"student_file": student_file, "teacher_file": teacher_file}


def _add_answer_key_to_doc(doc, questions, lang, as_teacher_cover=True):
    if as_teacher_cover:
        h = doc.add_heading(L("teacher_copy", lang), level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(L("answer_key", lang), level=1)
    for i, q in enumerate(questions, start=1):
        qtype = q.get("type")
        if qtype == "fill_blank":
            ans = q.get("answer", "—")
            if isinstance(ans, list):
                ans = ", ".join(ans)
        elif qtype == "match_column":
            left = q.get("left", [])
            right = q.get("right", [])
            ans = ", ".join(f"{l} → {r}" for l, r in zip(left, right))
        else:
            ans = q.get("answer", "—")
        p = doc.add_paragraph(f"{i}. {ans}")
        if is_urdu_text(str(ans)):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("")


def export_to_docx(config, questions, worksheet_title, grade, subject,
                    num_copies=1, randomize_mode="same",
                    lang="en", answer_key_mode="separate",
                    filepath="worksheet.docx"):
    doc = Document()
    last_copy_questions = None

    for copy_num in range(max(1, num_copies)):
        copy_questions = randomize_question_order(questions, randomize_mode)
        last_copy_questions = copy_questions

        if config.get("school_logo_base64"):
            try:
                header_img_data = base64.b64decode(config["school_logo_base64"].split(",")[1])
                img_stream = io.BytesIO(header_img_data)
                doc.add_picture(img_stream, width=Inches(1.3))
            except Exception:
                pass
        title = doc.add_heading(config.get("school_name", ""), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_heading(worksheet_title, level=1)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph(f"{L('grade', lang)}: {grade}   |   {L('subject', lang)}: {subject}")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = f"{L('name', lang)}: ____________"
        hdr[1].text = f"{L('class', lang)}: ______"
        hdr[2].text = f"{L('roll_no', lang)}: ______"
        hdr[3].text = f"{L('date', lang)}: ______"
        doc.add_paragraph("")

        for i, q in enumerate(copy_questions, start=1):
            qtext = q.get("text", "")
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {qtext}")
            run.bold = True
            if is_urdu_text(qtext):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            qtype = q.get("type")
            if qtype == "mcq":
                for j, opt in enumerate(q.get("options", [])):
                    doc.add_paragraph(f"   {chr(97+j)}) {opt}")
            elif qtype == "fill_blank":
                doc.add_paragraph("   " + "____________  " * q.get("blanks", 1))
            elif qtype == "match_column":
                left = q.get("left", [])
                if "_shuffled_right" not in q:
                    right = q.get("right", [])[:]
                    random.shuffle(right)
                    q["_shuffled_right"] = right
                right = q["_shuffled_right"]
                if left:
                    mt = doc.add_table(rows=len(left), cols=3)
                    for r, (l, rgt) in enumerate(zip(left, right)):
                        mt.rows[r].cells[0].text = l
                        mt.rows[r].cells[1].text = "—"
                        mt.rows[r].cells[2].text = rgt
            elif qtype in ("short", "detailed"):
                lines = q.get("lines", 3 if qtype == "short" else 6)
                for _ in range(lines):
                    doc.add_paragraph("_" * 60)
            doc.add_paragraph("")

        if answer_key_mode in ("inline", "both"):
            doc.add_page_break()
            _add_answer_key_to_doc(doc, copy_questions, lang, as_teacher_cover=False)

        if copy_num < num_copies - 1:
            doc.add_page_break()

    doc.save(filepath)

    teacher_filepath = None
    if answer_key_mode in ("separate", "both"):
        teacher_doc = Document()
        _add_answer_key_to_doc(teacher_doc, last_copy_questions, lang, as_teacher_cover=True)
        teacher_filepath = filepath.replace(".docx", "_ANSWER_KEY.docx")
        teacher_doc.save(teacher_filepath)

    return {"student_file": filepath, "teacher_file": teacher_filepath}


# ------------------------------------------------------------
# Free text -> questions parser (Method 3 style, no AI needed)
# ------------------------------------------------------------

def parse_questions_from_text(text):
    """
    Dependency-free parser for pasted text (Method 3 style, no AI needed).
    Detects: MCQ (with a/b/c/d options, resolves lettered answers to the
    actual option text), Fill in the Blank (___ runs), Match the Column
    ('Match the following' + Column A / Column B lists), and falls back to
    Short/Detailed for anything else. Looks for a trailing 'Answers:'
    section to fill in the answer key when present. Never invents an
    answer it can't find — it flags it instead.
    """
    if not text:
        return []

    q_start_pattern = re.compile(r'^\s*(\d+)[\.\)]\s*(.*)$')
    opt_pattern = re.compile(r'^\s*[\(\[]?([a-dA-D])[\)\.\]]\s*(.+)$')
    blank_pattern = re.compile(r'_{3,}')
    answers_header = re.compile(r'^\s*(answers?|answer\s*key)\s*[:：]?\s*$', re.IGNORECASE)
    answer_line_pattern = re.compile(r'^\s*(\d+)[\.\)]\s*(.+)$')
    match_header_pattern = re.compile(r'match\s+the\s+following', re.IGNORECASE)
    column_a_pattern = re.compile(r'^\s*column\s*a\s*:?\s*(.*)$', re.IGNORECASE)
    column_b_pattern = re.compile(r'^\s*column\s*b\s*:?\s*(.*)$', re.IGNORECASE)

    raw_lines = text.splitlines()

    # Split off a trailing "Answers:" section, if present
    split_idx = None
    for i, l in enumerate(raw_lines):
        if answers_header.match(l.strip()):
            split_idx = i
            break
    answers = {}
    if split_idx is not None:
        for al in raw_lines[split_idx + 1:]:
            m = answer_line_pattern.match(al.strip())
            if m:
                answers[int(m.group(1))] = m.group(2).strip()
        raw_lines = raw_lines[:split_idx]

    # Group lines into one block per numbered question
    blocks = []
    current_num, current_lines = None, []
    for l in raw_lines:
        m = q_start_pattern.match(l)
        if m:
            if current_lines:
                blocks.append((current_num, current_lines))
            current_num = int(m.group(1))
            current_lines = [m.group(2)]
        elif current_lines is not None and current_num is not None:
            current_lines.append(l)
    if current_lines and current_num is not None:
        blocks.append((current_num, current_lines))

    result = []
    for q_num, block_lines in blocks:
        lines = [l for l in block_lines if l.strip()]
        if not lines:
            continue
        header = lines[0].strip()
        body = lines[1:]

        # --- Match the Column ---
        if match_header_pattern.search(header):
            left, right = [], []
            for l in body:
                ma = column_a_pattern.match(l.strip())
                mb = column_b_pattern.match(l.strip())
                if ma and ma.group(1):
                    left = [x.strip() for x in ma.group(1).split(",") if x.strip()]
                elif mb and mb.group(1):
                    right = [x.strip() for x in mb.group(1).split(",") if x.strip()]
            if left and right and len(left) == len(right):
                result.append({"type": "match_column", "text": header, "left": left, "right": right})
                continue
            # falls through to short/detailed if columns weren't parseable

        # --- MCQ ---
        options = []
        for l in body:
            om = opt_pattern.match(l.strip())
            if om:
                options.append(om.group(2).strip())
        if len(options) >= 2:
            raw_ans = answers.get(q_num)
            answer_text = None
            if raw_ans:
                letter_m = re.match(r'^\(?([a-dA-D])\)?\.?$', raw_ans.strip())
                if letter_m:
                    idx = ord(letter_m.group(1).lower()) - ord('a')
                    if 0 <= idx < len(options):
                        answer_text = options[idx]
                else:
                    answer_text = raw_ans.strip()
            result.append({
                "type": "mcq", "text": header, "options": options,
                "answer": answer_text or "NOT PROVIDED",
            })
            continue

        # --- Fill in the Blank ---
        if blank_pattern.search(header):
            blanks = len(blank_pattern.findall(header))
            result.append({
                "type": "fill_blank", "text": header, "blanks": blanks,
                "answer": answers.get(q_num) or "NOT PROVIDED",
            })
            continue

        # --- Short / Detailed (fallback) ---
        full_text = " ".join([header] + body).strip()
        is_detailed = bool(re.search(r'\b(explain|describe|discuss)\b.*\b(detail|paragraph|essay)\b',
                                      full_text, re.IGNORECASE)) or len(full_text) > 140
        result.append({
            "type": "detailed" if is_detailed else "short",
            "text": full_text,
            "lines": 6 if is_detailed else 3,
            "answer": answers.get(q_num) or "NOT PROVIDED",
        })

    return result
